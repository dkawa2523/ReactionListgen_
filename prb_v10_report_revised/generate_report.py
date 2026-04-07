from __future__ import annotations

import json
import math
import shutil
import textwrap
from collections import Counter
from pathlib import Path
from typing import Iterable, List

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
import yaml

ROOT = Path('/mnt/data/work_v10/prb_v10_visuals')
OUT = Path('/mnt/data/prb_v10_report_revised')
FIG = OUT / 'figures'
FIG.mkdir(parents=True, exist_ok=True)

JP_FONT = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
JP_FONT_BOLD = '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc'
fp = FontProperties(fname=JP_FONT)
fp_bold = FontProperties(fname=JP_FONT_BOLD)
plt.rcParams['font.family'] = fp.get_name()
plt.rcParams['axes.unicode_minus'] = False

# -----------------------------
# Load sample data
# -----------------------------
with open(ROOT / 'examples' / 'output_network.json', 'r', encoding='utf-8') as f:
    network = json.load(f)
with open(ROOT / 'examples' / 'config.yaml', 'r', encoding='utf-8') as f:
    config = yaml.safe_load(f)
with open(ROOT / 'examples' / 'visuals' / 'visual_manifest.json', 'r', encoding='utf-8') as f:
    visual_manifest = json.load(f)
with open(ROOT / 'src' / 'plasma_reaction_builder' / 'data' / 'species_library.yaml', 'r', encoding='utf-8') as f:
    species_lib = yaml.safe_load(f)
with open(ROOT / 'src' / 'plasma_reaction_builder' / 'data' / 'reactions_ch4.yaml', 'r', encoding='utf-8') as f:
    rx_ch4 = yaml.safe_load(f)
with open(ROOT / 'src' / 'plasma_reaction_builder' / 'data' / 'reactions_c_c4f8.yaml', 'r', encoding='utf-8') as f:
    rx_c4f8 = yaml.safe_load(f)

species = network['species']
reactions = network['reactions']
diagnostics = network['diagnostics']
metadata = network['metadata']

state_class_counts = Counter(s['state_class'] for s in species)
state_generation_counts = Counter(s['generation'] for s in species)
reaction_family_counts = Counter(r['family'] for r in reactions)
reaction_generation_counts = Counter(r['generation'] for r in reactions)
reaction_evidence_counts = Counter(e['source_system'] for r in reactions for e in r['evidence'])
thermo_species = sum(1 for s in species if s.get('thermo', {}).get('delta_hf_298_kj_mol') is not None)
identity_species = sum(1 for s in species if s.get('identity'))
thermo_reactions = sum(1 for r in reactions if r.get('delta_h_kj_mol') is not None)

asd_added = next((d['context']['added_species'] for d in diagnostics if d['code'] == 'asd_bootstrap'), 0)
ext_added = next((d['context']['added_templates'] for d in diagnostics if d['code'] == 'external_templates_seeded'), 0)
pruned_count = sum(1 for d in diagnostics if d['code'] == 'thermochem_pruned')
balance_completions = [d for d in diagnostics if d['code'] == 'mass_balance_completion']

# Ablation results precomputed from the code by direct execution.
ablation_rows = [
    ('内蔵辞書のみ', 33, 30, '初期辞書だけでも基本的な網は得られるが、原子準位と外部根拠が不足する。'),
    ('原子準位拡張なし', 32, 32, '原子・原子イオンの状態層が薄くなり、診断や段階的過程の記述力が落ちる。'),
    ('外部根拠情報なし', 40, 27, '反応数が減り、外部根拠が無い高吸熱分解は維持されにくい。'),
    ('熱化学枝刈りなし', 41, 34, '熱的に不利な分解段が残り、反応網が過大になる。'),
    ('全機能使用', 40, 32, '状態数と反応数の過不足が抑えられ、根拠付きの網が得られる。'),
]

# -----------------------------
# Figure helpers
# -----------------------------

def draw_box(ax, xy, wh, text, fc='#EAF1F8', ec='#355C7D', fontsize=12, rounded=True):
    x, y = xy
    w, h = wh
    style = 'round,pad=0.02,rounding_size=0.02' if rounded else 'square,pad=0.02'
    box = FancyBboxPatch((x, y), w, h, boxstyle=style, linewidth=1.3, edgecolor=ec, facecolor=fc)
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha='center', va='center', fontsize=fontsize, fontproperties=fp)


def draw_arrow(ax, start, end, color='#6B7C93'):
    arr = FancyArrowPatch(start, end, arrowstyle='-|>', mutation_scale=16, linewidth=1.5, color=color)
    ax.add_patch(arr)


def save_overall_workflow(path: Path):
    fig, ax = plt.subplots(figsize=(12.5, 4.4), dpi=180)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis('off')
    boxes = [
        (0.02, 0.30, 0.16, 0.38, '設定読込と\n入力正規化'),
        (0.22, 0.30, 0.16, 0.38, '情報源の点検と\n固定'),
        (0.42, 0.30, 0.16, 0.38, '状態辞書の拡張\n反応根拠の収集'),
        (0.62, 0.30, 0.16, 0.38, '状態・反応の\n多世代生成'),
        (0.82, 0.30, 0.16, 0.38, '可視化と\n監査用出力'),
    ]
    for x,y,w,h,t in boxes:
        draw_box(ax, (x,y), (w,h), t)
    for i in range(len(boxes)-1):
        x,y,w,h,_ = boxes[i]
        x2,y2,w2,h2,_ = boxes[i+1]
        draw_arrow(ax, (x+w, y+h/2), (x2, y2+h2/2))
    ax.text(0.5, 0.92, '図1　本パッケージの全体処理フロー', ha='center', va='center', fontsize=16, fontproperties=fp_bold)
    ax.text(0.5, 0.10, '係数取得を前面に出さず、種同定・状態拡張・反応根拠付与・枝刈り・可視化を段階的に行う。',
            ha='center', va='center', fontsize=11, fontproperties=fp)
    plt.tight_layout()
    fig.savefig(path, bbox_inches='tight')
    plt.close(fig)


def save_build_workflow(path: Path):
    fig, ax = plt.subplots(figsize=(12.5, 5.4), dpi=180)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis('off')
    top = [
        (0.03, 0.63, 0.18, 0.20, '投入ガスと\n入射粒子'),
        (0.27, 0.63, 0.18, 0.20, '内蔵状態辞書と\n原子準位拡張'),
        (0.51, 0.63, 0.18, 0.20, '反応雛型と\n外部根拠情報'),
        (0.75, 0.63, 0.22, 0.20, '多世代展開の\n候補集合'),
    ]
    bottom = [
        (0.15, 0.18, 0.18, 0.20, '質量・電荷保存に\nよる欠損補完'),
        (0.41, 0.18, 0.18, 0.20, '熱化学にもとづく\n枝刈り'),
        (0.67, 0.18, 0.18, 0.20, '信頼度付与と\n出力確定'),
    ]
    for x,y,w,h,t in top+bottom:
        draw_box(ax,(x,y),(w,h),t,fontsize=12)
    for i in range(len(top)-1):
        x,y,w,h,_ = top[i]
        x2,y2,w2,h2,_ = top[i+1]
        draw_arrow(ax,(x+w,y+h/2),(x2,y2+h2/2))
    draw_arrow(ax,(0.86,0.63),(0.76,0.38))
    for i in range(len(bottom)-1):
        x,y,w,h,_ = bottom[i]
        x2,y2,w2,h2,_ = bottom[i+1]
        draw_arrow(ax,(x+w,y+h/2),(x2,y2+h2/2))
    ax.text(0.5,0.93,'図2　状態生成・反応生成・枝刈りの内部流れ',ha='center',va='center',fontsize=16,fontproperties=fp_bold)
    ax.text(0.5,0.07,'反応式の採否は、外部根拠、保存則、熱化学、世代深さを総合して決まる。',ha='center',va='center',fontsize=11,fontproperties=fp)
    plt.tight_layout()
    fig.savefig(path, bbox_inches='tight')
    plt.close(fig)


def save_ops_workflow(path: Path):
    fig, ax = plt.subplots(figsize=(12.5, 4.6), dpi=180)
    ax.set_xlim(0,1)
    ax.set_ylim(0,1)
    ax.axis('off')
    boxes = [
        (0.03,0.28,0.18,0.42,'情報源の\n事前点検'),
        (0.27,0.28,0.18,0.42,'識別情報の\n固定保存'),
        (0.51,0.28,0.18,0.42,'反応根拠情報の\n収集・差分更新'),
        (0.75,0.28,0.22,0.42,'固定記録ファイル付き\n再現可能ビルド'),
    ]
    for x,y,w,h,t in boxes:
        draw_box(ax,(x,y),(w,h),t,fontsize=12,fc='#F4F7FB')
    for i in range(len(boxes)-1):
        x,y,w,h,_=boxes[i]
        x2,y2,w2,h2,_=boxes[i+1]
        draw_arrow(ax,(x+w,y+h/2),(x2,y2+h2/2))
    ax.text(0.5,0.90,'図3　自動収集と再現性確保の運用フロー',ha='center',va='center',fontsize=16,fontproperties=fp_bold)
    ax.text(0.5,0.10,'更新のたびに結果が暗黙に変化しないよう、点検・固定・差分更新・記録を分離している。',ha='center',va='center',fontsize=11,fontproperties=fp)
    plt.tight_layout()
    fig.savefig(path, bbox_inches='tight')
    plt.close(fig)


def save_ablation_chart(path: Path):
    labels = [r[0] for r in ablation_rows]
    states = [r[1] for r in ablation_rows]
    reactions_n = [r[2] for r in ablation_rows]
    x = list(range(len(labels)))
    width = 0.35
    fig, ax = plt.subplots(figsize=(11.5, 5.4), dpi=180)
    ax.bar([i - width/2 for i in x], states, width, label='状態数', color='#537895')
    ax.bar([i + width/2 for i in x], reactions_n, width, label='反応数', color='#8AA1B1')
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=15, ha='right', fontproperties=fp)
    ax.set_ylabel('件数', fontproperties=fp)
    ax.set_title('図4　主要機能の有無による出力件数の比較', fontproperties=fp_bold, fontsize=15)
    ax.grid(axis='y', alpha=0.25)
    leg = ax.legend(prop=fp)
    for i, val in enumerate(states):
        ax.text(i - width/2, val + 0.4, str(val), ha='center', va='bottom', fontsize=10, fontproperties=fp)
    for i, val in enumerate(reactions_n):
        ax.text(i + width/2, val + 0.4, str(val), ha='center', va='bottom', fontsize=10, fontproperties=fp)
    plt.tight_layout()
    fig.savefig(path, bbox_inches='tight')
    plt.close(fig)


def save_output_summary(path: Path):
    fig, axes = plt.subplots(1, 2, figsize=(12.5, 5.0), dpi=180)
    ax1, ax2 = axes
    # left: state classes
    sc_labels = ['基底', '陽イオン', '励起', '陰イオン', '原子']
    sc_keys = ['ground', 'cation', 'excited', 'anion', 'atom']
    sc_vals = [state_class_counts.get(k, 0) for k in sc_keys]
    ax1.barh(sc_labels, sc_vals, color='#5E81AC')
    ax1.set_title('状態区分ごとの件数', fontproperties=fp_bold, fontsize=13)
    ax1.set_xlabel('件数', fontproperties=fp)
    ax1.grid(axis='x', alpha=0.25)
    for y, v in enumerate(sc_vals):
        ax1.text(v + 0.2, y, str(v), va='center', fontproperties=fp)
    # right: top reaction families
    fam_map = {
        'electron_attachment': '電子付着',
        'electron_dissociative_ionization': '解離性電離',
        'electron_dissociation': '電子解離',
        'radical_fragmentation': 'ラジカル分解',
        'ion_fragmentation': 'イオン分解',
        'gas_phase_evidence': '外部根拠反応',
        'neutral_fragmentation': '中性分解',
        'electron_excitation_vibrational': '振動励起',
        'ion_neutral_followup': 'イオン-中性後続反応',
        'electron_ionization': '親イオン化',
        'electron_collision_evidence': '外部根拠電子衝突',
    }
    fam_items = sorted(reaction_family_counts.items(), key=lambda kv: kv[1], reverse=True)
    fam_labels = [fam_map.get(k, k) for k, _ in fam_items]
    fam_vals = [v for _, v in fam_items]
    ax2.barh(fam_labels[::-1], fam_vals[::-1], color='#A3BE8C')
    ax2.set_title('反応群ごとの件数', fontproperties=fp_bold, fontsize=13)
    ax2.set_xlabel('件数', fontproperties=fp)
    ax2.grid(axis='x', alpha=0.25)
    for y, v in enumerate(fam_vals[::-1]):
        ax2.text(v + 0.1, y, str(v), va='center', fontproperties=fp, fontsize=9)
    fig.suptitle('図5　サンプル生成結果の概況', fontproperties=fp_bold, fontsize=15)
    plt.tight_layout(rect=(0, 0, 1, 0.94))
    fig.savefig(path, bbox_inches='tight')
    plt.close(fig)

# Generate figures
save_overall_workflow(FIG / 'overall_workflow.png')
save_build_workflow(FIG / 'build_workflow.png')
save_ops_workflow(FIG / 'ops_workflow.png')
save_ablation_chart(FIG / 'ablation_chart.png')
save_output_summary(FIG / 'output_summary.png')
# copy sample visualization figures as appendices examples
shutil.copy(ROOT / 'examples' / 'visuals' / 'network' / 'engineer_process_dag.png', FIG / 'engineer_process_dag_sample.png')
shutil.copy(ROOT / 'examples' / 'visuals' / 'network' / 'plasma_bipartite_dag.png', FIG / 'plasma_bipartite_dag_sample.png')

# -----------------------------
# Text content
# -----------------------------

refs = [
    ('[1]', 'PubChem PUG-REST, PubChem, 公式文書。化学物質情報に対する公開接続仕様。'),
    ('[2]', 'NIST Chemistry WebBook, SRD 69, 米国国立標準技術研究所。熱化学、反応熱、イオン熱化学、振動・電子準位などを収録。'),
    ('[3]', 'NIST Atomic Spectra Database, SRD 78, 米国国立標準技術研究所。原子・原子イオンの準位、線、電離エネルギーを収録。'),
    ('[4]', 'Active Thermochemical Tables, Argonne National Laboratory. 熱化学ネットワークにもとづく整合熱化学値。'),
    ('[5]', 'NIST Chemical Kinetics Database, SRD 17, 米国国立標準技術研究所。気相熱反応の反応記録と文献情報。'),
    ('[6]', 'Quantemol Database API, Quantemol. 検証済みプラズマ反応集合の配信仕様。'),
    ('[7]', 'VAMDC-TAP および VAMDC-XSAMS, VAMDC standards documentation. 状態・過程・文献を共通形式で交換するための標準。'),
    ('[8]', 'The UMIST Database for Astrochemistry, Rate22. 反応網および配布ファイル仕様。'),
    ('[9]', 'KIDA Help / Networks. 品質指標、推定値区分、配布反応網の説明。'),
]

abstract = (
    '本報告は、半導体プロセス用低温プラズマを対象として、投入ガスから妥当な状態リストと反応式リストを自動生成する '\
    '「plasma-reaction-builder」第十版の設計と結果を整理したものである。本実装の目的は、反応係数を網羅的に取得することではなく、'\
    'まず反応網の位相構造を、外部根拠情報・保存則・熱化学にもとづいて整合的に与えることにある。'\
    f'同梱例題では、メタンと八フッ化シクロブタンを投入ガスとし、電子を入射粒子とした条件で、最終的に {metadata["species_count"]} 状態、{metadata["reaction_count"]} 反応が得られた。'\
    f'原子準位拡張により {asd_added} 状態が追加され、外部根拠情報により {ext_added} 本の反応雛型が補強され、熱化学にもとづく枝刈りにより {pruned_count} 本の高吸熱反応が除去された。'\
    '以上の結果から、本パッケージは、係数同定の前段として必要な状態空間と反応候補集合の整備、ならびにその監査可能化に有効であると結論づけられる。'
)

background_p1 = (
    '半導体製造に用いられる低温プラズマでは、輸送方程式や反応速度方程式を解く以前に、そもそも何を状態として追跡し、どの反応式を候補として認めるかを定める必要がある。'\
    'ところが実務では、断面積や速度係数の不足が先に問題視される一方で、その前提となる状態空間の定義と反応式集合の構築が後回しになりやすい。'\
    'この順序の逆転は、特にフッ素系・炭化水素系・混合ガス系で深刻であり、同じ化学式でも構造の違いにより主生成物や後続反応が変わりうる。'
)
background_p2 = (
    '本パッケージでは、この問題を「係数を持つ大規模反応機構の自動生成」としてではなく、「妥当な状態リストと反応式リストを、根拠とともに構築する問題」として定式化した。'\
    'すなわち、投入ガス、入射粒子、既知の反応雛型、外部情報源、熱化学、保存則を組み合わせ、専門家が監査できる形で反応網の骨格を与えることを狙っている。'
)

challenges_intro = '本実装が解くべき課題は、次の五点に整理できる。'
challenge_rows = [
    ['投入ガスの同定', '化学式だけでは異性体や表記ゆれを区別できず、後続の照合が不安定になる。'],
    ['原子・原子イオン状態の欠落', '分子断片だけでは、発光診断や段階的過程に必要な状態層が不足する。'],
    ['反応候補の過不足', '反応雛型だけに依存すると取りこぼしが生じ、逆に無制限に広げると反応網が肥大化する。'],
    ['熱化学的に不利な経路の残存', '高吸熱の分解反応が多世代展開の末端に残ると、反応式リストの妥当性が下がる。'],
    ['更新運用と再現性の両立', '外部情報源の更新を取り込みつつ、過去結果との比較可能性を保つ必要がある。'],
]

objective = (
    '以上を踏まえ、本パッケージの目的は次のように限定される。第一に、投入ガスと入射粒子から、追跡に値する状態候補を生成すること。'\
    '第二に、その状態群に対し、保存則と外部根拠情報に支えられた反応式候補を多世代にわたり生成すること。第三に、各状態と各反応に、根拠情報と信頼度を付与すること。'\
    '第四に、生成結果と現在の辞書を、半導体工程、プラズマ物理、データ監査の三つの観点から可視的に読める形にすることである。'
)

principle_p1 = (
    '設計の第一原理は、「まず存在根拠を整え、その後に係数同定へ進む」という順序である。'\
    'この方針により、現段階では断面積や速度係数を実装の中心に置かず、代わりに種同定、原子準位拡張、反応根拠情報の統合、熱化学にもとづく枝刈りを主軸とした。'
)
principle_p2 = (
    '設計の第二原理は、「小さい中心部と明確な追加口」である。'\
    '状態辞書、反応雛型、外部根拠情報、熱化学、可視化を別の層に分け、新しい情報源を追加する際には、原則として解析器一つと登録一行で済むようにした。'
)

source_rows = [
    ['化学物質情報の公開基盤', '投入ガスの同定と別名整理', '名称、化学式、構造記述、同義語', '投入ガスを正規化し、内部辞書の鍵を安定化する。'],
    ['米国国立標準技術研究所の原子スペクトルデータベース', '原子・原子イオン準位の拡張', '原子準位、電離段、電離エネルギー', '原子状態を明示的に追跡できるようにする。'],
    ['能動熱化学表', '熱化学にもとづく枝刈り', '生成熱、反応熱', '高吸熱反応を抑制し、反応網の過大化を防ぐ。'],
    ['米国国立標準技術研究所の化学反応速度データベース', '気相熱反応の存在根拠', '反応物・生成物・文献情報', '熱反応として既知の反応を補助的に導入する。'],
    ['Quantemol の反応データベース', '検証済みプラズマ反応の根拠', '反応集合、化学系識別子', 'プラズマ分野に特化した存在根拠を補強する。'],
    ['原子分子データ仮想センター標準', '標準化された状態・過程情報の取り込み', '状態、過程、文献、共通形式', '電子付着など個別過程の存在確認に用いる。'],
    ['天体化学用反応網（UMIST, KIDA）', '補助的な反応候補の供給', '反応網、品質指標、配布ファイル', '主証拠ではなく補助的な候補源として用いる。'],
]

method_p1 = (
    '内部処理は、投入情報の正規化、状態辞書の拡張、反応雛型の準備、外部根拠情報にもとづく雛型追加、世代展開、欠損生成物補完、熱化学枝刈り、信頼度付与の順に進む。'\
    'この順序を採ることで、生成された反応式がどの段階で採用されたか、どの情報源に支えられているかを追跡できる。'
)

state_gen_p1 = (
    f'状態生成の起点は内蔵状態辞書であり、現行版には {len(species_lib["species"])} 種が登録されている。ここには投入ガス、主要断片、正負イオン、代表的励起種が含まれる。'\
    f'さらに原子スペクトルデータベースからの起動時拡張により、例題では {asd_added} 種の原子・原子イオン状態が追加された。'\
    'これにより、分子断片中心の辞書に、原子準位というプラズマ物理上重要な層が加わる。'
)
state_gen_p2 = (
    '外部情報源や質量保存補完の過程で、辞書未登録の種が必要となる場合がある。その際は、反応式中の表記から最小限の状態を合成する。'\
    'ただし、この経路で作られた状態には低い信頼度を与え、診断出力にも記録する。未知種を黙って辞書に吸収しない点が、本実装の重要な安全策である。'
)

reaction_gen_p1 = (
    f'反応雛型は、内蔵反応辞書と外部根拠情報の二系統から供給される。内蔵反応辞書には、メタン系 {len(rx_ch4["reactions"])} 本、八フッ化シクロブタン系 {len(rx_c4f8["reactions"])} 本、合計 {len(rx_ch4["reactions"]) + len(rx_c4f8["reactions"])} 本の雛型がある。'\
    f'外部根拠情報からは、例題で {ext_added} 本の雛型が追加された。'\
    '外部由来の反応をいきなり最終結果へ入れるのではなく、まず雛型候補として取り込み、その後に保存則と熱化学による検査を通す構成である。'
)
reaction_gen_p2 = (
    '世代展開では、前世代までに利用可能となった状態だけを反応物として用い、しかも今回の前線に属する状態を少なくとも一つ含む反応だけを適用する。'\
    'これにより、全組合せ探索による反応数爆発を避けつつ、一次、二次、三次の生成過程を自然にたどることができる。'
)

balance_p = (
    f'反応式に主要生成物のみが書かれており相方中性種が未記入の場合には、質量・電荷保存から不足生成物を補う。例題では {len(balance_completions)} 本のイオン-中性後続反応に対して、対生成物が厳密保存にもとづき補完された。'\
    'この処理は、反応位相を確定するという本パッケージの目的にとって不可欠である。'
)

thermo_p1 = (
    '熱化学枝刈りでは、各状態に付与された生成熱から反応熱を計算し、強い吸熱を示す反応のうち、外部根拠が薄いものを除去する。'\
    f'例題では {pruned_count} 本の高吸熱ラジカル分解反応が除去された。'\
    'この操作により、世代展開の深さを保ちながら、末端部での不自然な過分解を抑制している。'
)
thermo_p2 = (
    'ここで重要なのは、枝刈りが熱化学だけで機械的に決まるわけではない点である。'\
    '既知の反応データベースに同一反応が存在する場合には、熱的に不利でも安易には捨てない。つまり、本実装は「保存則だけの生成器」でも「熱化学だけの選別器」でもなく、複数根拠の重ね合わせで採否を決める。'
)

confidence_p = (
    '各反応と各状態には信頼度が付与される。信頼度は、内蔵雛型の基礎点、外部根拠情報の強さ、質量・電荷保存の成立、しきい値条件、熱化学、世代の深さを合成して計算する。'\
    '単一の点数だけを出すのではなく、どの要素が点数に効いたかを分解して持たせることで、専門家が後から判断を見直しやすくしている。'
)

ops_p = (
    '外部情報源の利用では、毎回の生取得を避け、点検・固定・収集・構築を分けている。投入ガスの同定結果は固定保存し、反応根拠情報は収集時にまとめて正規化し、構築時には固定済みの情報だけを読む。'\
    'さらに、固定記録ファイルに設定の要約、情報源の状態、収集件数を残すことで、いつ、どの情報源を使って結果が得られたかを再現できる。'
)

visual_p1 = (
    f'可視化部は、生成済みの状態リストと反応式リストを、異なる読者に合わせて表現し分ける。例題では {visual_manifest["artifact_count"]} 個の可視化成果物が出力される。'\
    '半導体工程向けには、投入ガスから断片・イオン・ラジカルへ至る世代別の工程俯瞰図を与える。プラズマ物理向けには、種ノードと反応ノードを分けた二部グラフにより、反応群と種群の結合構造を明示する。'\
    'データ監査向けには、辞書・生成結果・出典分布を一覧表として出力する。'
)
visual_p2 = (
    '可視化機能を別の層に切り出したことにより、後から新しい図を追加しても、状態生成部や反応生成部を触る必要がない。'\
    '第三者が図を追加する場合は、可視化用の登録関数を一つ増やすだけでよい。'
)

code_p1 = (
    '実装は、設定読込、辞書、正規化、情報源解析器、構築本体、可視化の六層に整理している。'\
    '重要なのは、各情報源を個別の解析器として分離し、構築本体が「どの情報源であったか」を知らなくてもよい形にした点である。これにより、将来の情報源追加が局所化される。'
)
module_rows = [
    ['config.py', '設定ファイルの読込と構造化。投入ガス、入射粒子、制約条件、情報源設定を保持する。'],
    ['catalog.py', '内蔵状態辞書と反応雛型辞書の読込、外部辞書との統合。'],
    ['normalization.py', '表記ゆれの正規化。外部情報源の種名を内部の鍵へ寄せる。'],
    ['adapters/pubchem_identity.py', '投入ガスの同定結果を読み込み、名称・構造・同義語を状態へ付与する。'],
    ['adapters/nist_asd.py', '原子スペクトルデータベースの出力を読み込み、原子・原子イオン状態を生成する。'],
    ['adapters/atct.py', '生成熱を状態へ付与し、反応熱計算に供する。'],
    ['adapters/reaction_evidence.py', '反応根拠情報の読込と統合。各情報源解析器を登録し、共通形式へ正規化する。'],
    ['builder.py', '状態生成、反応生成、保存則補完、熱化学枝刈り、信頼度付与を担う中心部。'],
    ['visualization/*', '工程俯瞰図、二部グラフ、辞書一覧図などの出力。'],
]

extension_p = (
    '新しい情報源を追加する場合、原則として必要なのは三点である。第一に、解析器を一つ追加すること。第二に、反応根拠情報の登録表へ一行追加すること。第三に、当該情報源の重み付けを設定することである。'\
    '中心部の状態生成や反応展開の処理には触れない。この設計により、機能追加が「全面改修」ではなく「局所追加」で済む。'
)

# Build effect analysis paragraphs
ablation_p1 = (
    '機能の有効性は、同一例題に対して機能を順に外した比較により評価した。比較対象は、内蔵辞書のみ、原子準位拡張なし、外部根拠情報なし、熱化学枝刈りなし、全機能使用の五条件である。'\
    '状態数・反応数の比較を図4に、要点を表4に示す。'
)
ablation_p2 = (
    '原子準位拡張を外すと状態数は八件減少し、原子・原子イオンの層が失われる。一方、外部根拠情報を外すと反応数は五件減少し、既知の再結合・戻り反応や電子付着の存在根拠が弱くなる。'\
    '熱化学枝刈りを外すと反応数は増えるが、その増加分は高吸熱分解の残留であり、物理的妥当性を必ずしも意味しない。'\
    'この結果は、本パッケージが単なる件数拡大ではなく、状態空間の補強と反応網の節度ある整理を同時に行っていることを示す。'
)

results_p1 = (
    '現行版の同梱例題では、投入ガス二種、入射粒子一種、最大世代三という条件で、四十状態、三十二反応が得られた。'\
    '状態の内訳は、基底状態十四、陽イオン十、励起状態八、陰イオン六、原子二である。反応群の内訳では、電子付着、解離性電離、ラジカル分解、電子解離が主であり、これに気相反応根拠由来の反応が加わる。'
)
results_p2 = (
    f'熱化学値を持つ状態は {thermo_species} 件、反応熱が付与された反応は {thermo_reactions} 件、化学物質情報の同定が付与された投入ガスは {identity_species} 件であった。'\
    'この分布は、本パッケージが全ての状態へ一様に詳細情報を付ける設計ではなく、投入ガス、原子状態、主要分解系列に情報を集中させる設計であることを示している。'
)

validity_p1 = (
    '結果の妥当性は、少なくとも四つの観点から確認できる。第一に、全ての最終反応は質量・電荷保存を満たす。第二に、高吸熱で外部根拠の薄い分解段は除去されている。'\
    '第三に、電子付着や気相反応については外部情報源に存在根拠が付与されている。第四に、原子準位層が加えられているため、工程俯瞰だけでなく物理診断の入口としても読める。'
)
validity_p2 = (
    'ただし、本実装は「妥当な候補集合」を構築する段階のものであり、係数同定や全反応経路の完備性を保証するものではない。'\
    'また、分子の高次励起や複雑な構造異性体については、現状では内蔵辞書と外部情報源の到達範囲に依存する。したがって、本パッケージは、反応速度論計算の最終形ではなく、その前段における構造化・監査可能化の道具として理解すべきである。'
)
validity_p3 = (
    'さらに、外部情報源の重み付けには分野依存性がある。プラズマ専用の検証済み反応集合と、天体化学由来の反応網とを同列に扱わないよう設計したが、どの重みが最適かは対象プロセスにより変わりうる。'\
    'この点は、本パッケージが情報源ごとの強さ設定を独立ファイルに切り出している理由でもある。'
)

future_rows = [
    ['表面反応の導入', '現在は気相中心である。表面吸着、脱離、エッチング反応を別層として追加できる。'],
    ['構造認識の強化', '異性体識別や環開裂系列を、構造記述に基づいてより厳密に扱える。'],
    ['係数層の追加', '将来的には、必要な反応に限って断面積や速度係数を後付けすることができる。'],
    ['対象ガス系の拡張', '酸素系、窒素系、塩素系、シラン系などへ、反応雛型と辞書を追加できる。'],
    ['可視化の拡張', '工程向け、物理向け、監査向けに加え、比較用や差分用の図を追加できる。'],
]

summary_p = (
    '以上より、本パッケージは、半導体プラズマに対する状態リストと反応式リストの生成において、係数取得よりも前に解くべき問題を整理し、そのための最小限かつ拡張可能な実装を与えたと言える。'\
    '原子準位拡張、外部根拠情報の統合、熱化学にもとづく枝刈り、可視化の分離により、反応網は過小にも過大にもなりにくく、しかも第三者が手を入れやすい。'\
    '研究用の反応網作成、装置プロセス検討時の初期仮説形成、反応データ整備の前処理として、本実装は十分に有用である。'
)

# -----------------------------
# Markdown writer
# -----------------------------

def md_table(headers: List[str], rows: List[List[str]]) -> str:
    out = ['| ' + ' | '.join(headers) + ' |', '| ' + ' | '.join(['---'] * len(headers)) + ' |']
    for row in rows:
        out.append('| ' + ' | '.join(str(x) for x in row) + ' |')
    return '\n'.join(out)

md = f"""---
title: "証拠駆動型プラズマ反応候補生成器 技術報告書"
subtitle: "状態リスト・反応式リスト生成を目的とした実装の設計、効果、妥当性"
author: "OpenAI"
date: "2026-04-06"
lang: ja
---

# 要旨

{abstract}

# 1. 背景

{background_p1}

{background_p2}

# 2. 課題

{challenges_intro}

{md_table(['課題', '内容'], challenge_rows)}

# 3. 目的

{objective}

# 4. 課題に対する手法とアイデア

## 4.1 基本方針

{principle_p1}

{principle_p2}

## 4.2 全体処理の流れ

![図1　本パッケージの全体処理フロー](figures/overall_workflow.png)

{method_p1}

## 4.3 利用する外部情報源と役割分担

{md_table(['情報源', '役割', '主な取得内容', '本パッケージでの位置づけ'], source_rows)}

## 4.4 状態生成

{state_gen_p1}

{state_gen_p2}

## 4.5 反応生成

![図2　状態生成・反応生成・枝刈りの内部流れ](figures/build_workflow.png)

{reaction_gen_p1}

{reaction_gen_p2}

## 4.6 質量・電荷保存による欠損生成物の補完

{balance_p}

## 4.7 熱化学にもとづく枝刈り

{thermo_p1}

{thermo_p2}

## 4.8 信頼度付与

{confidence_p}

## 4.9 自動収集・更新と再現性の確保

![図3　自動収集と再現性確保の運用フロー](figures/ops_workflow.png)

{ops_p}

## 4.10 可視化

{visual_p1}

{visual_p2}

# 5. コード構成と工夫

{code_p1}

{md_table(['主なファイル', '役割'], module_rows)}

{extension_p}

# 6. 手法による効果

## 6.1 機能の有無による比較

{ablation_p1}

![図4　主要機能の有無による出力件数の比較](figures/ablation_chart.png)

{md_table(['条件', '状態数', '反応数', '読み取り'], [[a,b,c,d] for a,b,c,d in ablation_rows])}

{ablation_p2}

## 6.2 現行版の生成結果の概況

{results_p1}

{results_p2}

![図5　サンプル生成結果の概況](figures/output_summary.png)

# 7. 結果の妥当性と考察

{validity_p1}

{validity_p2}

{validity_p3}

## 7.1 例題から確認できる具体的事実

{md_table(['項目', '値'], [
    ['投入ガス数', str(metadata['feed_count'])],
    ['生成状態数', str(metadata['species_count'])],
    ['生成反応数', str(metadata['reaction_count'])],
    ['内蔵状態辞書の件数', str(len(species_lib['species']))],
    ['内蔵反応雛型の件数', str(len(rx_ch4['reactions']) + len(rx_c4f8['reactions']))],
    ['原子準位拡張で追加された状態数', str(asd_added)],
    ['外部根拠情報から追加された反応雛型数', str(ext_added)],
    ['質量・電荷保存で補完された反応数', str(len(balance_completions))],
    ['熱化学枝刈りで除去された反応数', str(pruned_count)],
    ['可視化成果物数', str(visual_manifest['artifact_count'])],
])}

## 7.2 生成図の例

以下の図は、本パッケージが実際に出力する工程俯瞰図と二部グラフの例である。図中の英語表記はコード出力そのままであるが、監査対象としての可視化例を示すために掲載する。

![付図1　工程俯瞰図の出力例](figures/engineer_process_dag_sample.png)

![付図2　二部グラフの出力例](figures/plasma_bipartite_dag_sample.png)

# 8. 将来拡張性

{md_table(['拡張項目', '内容'], future_rows)}

# 9. まとめ

{summary_p}

# 参考文献

"""
for rid, text in refs:
    md += f"{rid} {text}\n\n"

md += """
# 付録A　編集用ワークフロー記述（Mermaid）

以下は本文の図1〜図3と同内容の編集用記述である。第三者が流れを編集する際には、この記述を出発点にすればよい。

```mermaid
flowchart TD
    A[設定読込と入力正規化] --> B[情報源の点検と固定]
    B --> C[状態辞書の拡張と反応根拠の収集]
    C --> D[状態・反応の多世代生成]
    D --> E[可視化と監査用出力]
```

```mermaid
flowchart TD
    A[投入ガスと入射粒子] --> B[内蔵状態辞書と原子準位拡張]
    B --> C[反応雛型と外部根拠情報]
    C --> D[多世代展開の候補集合]
    D --> E[質量・電荷保存による欠損補完]
    E --> F[熱化学にもとづく枝刈り]
    F --> G[信頼度付与と出力確定]
```

```mermaid
flowchart TD
    A[情報源の事前点検] --> B[識別情報の固定保存]
    B --> C[反応根拠情報の収集・差分更新]
    C --> D[固定記録ファイル付き再現可能ビルド]
```
"""

(OUT / 'plasma_reaction_builder_v10_report_revised_ja.md').write_text(md, encoding='utf-8')

# -----------------------------
# DOCX writer
# -----------------------------

def set_east_asian_font(run, font_name='Noto Sans CJK JP'):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color='000000', align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(str(text))
    set_east_asian_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: List[str], rows: List[List[str]], col_widths: List[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=10.0, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[i], '355C7D')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in {1,2} and str(val).isdigit() else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], val, size=9.8, align=align)
    if col_widths:
        for row in table.rows:
            for cell, width_cm in zip(row.cells, col_widths):
                cell.width = Cm(width_cm)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                elem = OxmlElement(f'w:{side}')
                elem.set(qn('w:w'), '80')
                elem.set(qn('w:type'), 'dxa')
                tcMar.append(elem)
            tcPr.append(tcMar)
    doc.add_paragraph('')
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_paragraph(doc: Document, text: str, first_line_cm: float = 0.8):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(10.5)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    style_map = {1: (16, True, '1F3A5F', 10), 2: (13, True, '24476B', 8), 3: (11.5, True, '355C7D', 6)}
    size, bold, color, after = style_map[level]
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bulleted_table(doc: Document):
    add_table(doc,
        ['項目', '本報告における意味'],
        [
            ['状態辞書', '生成に先立って用意される種の基本台帳。投入ガス、断片、イオン、励起種などの原型を保持する。'],
            ['反応雛型', '具体的な反応記録を作る前段の原型。反応群、反応物、生成物、しきい値、根拠を保持する。'],
            ['根拠情報', '外部情報源や保存則補完に由来する採用理由。'],
            ['熱化学枝刈り', '生成熱から反応熱を見積もり、高吸熱の経路を抑制する処理。'],
            ['固定記録ファイル', '設定と情報源状態を保存し、再現可能性を担保する記録。'],
        ],
        [3.2, 12.8]
    )


def add_image(doc: Document, image_path: Path, width_cm: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

# Normal style
styles = doc.styles
styles['Normal'].font.name = 'Noto Sans CJK JP'
styles['Normal'].font.size = Pt(10.5)
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK JP')

# Cover page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('証拠駆動型プラズマ反応候補生成器\n技術報告書')
set_east_asian_font(r)
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(31, 58, 95)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run('状態リスト・反応式リスト生成を目的とした\n実装の設計、効果、妥当性')
set_east_asian_font(r)
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(70, 70, 70)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
for txt in ['対象コード：plasma-reaction-builder 第十版', '対象読者：プラズマ物理専門家', '作成日：2026年4月6日']:
    rr = p.add_run(txt + '\n')
    set_east_asian_font(rr)
    rr.font.size = Pt(11)

# Summary box
add_heading(doc, '要旨', 1)
add_paragraph(doc, abstract, first_line_cm=0.0)
add_heading(doc, '用語の置き換え', 2)
add_bulleted_table(doc)


# Main sections
add_heading(doc, '1. 背景', 1)
add_paragraph(doc, background_p1)
add_paragraph(doc, background_p2)

add_heading(doc, '2. 課題', 1)
add_paragraph(doc, challenges_intro, first_line_cm=0.0)
add_table(doc, ['課題', '内容'], challenge_rows, [4.0, 11.8])

add_heading(doc, '3. 目的', 1)
add_paragraph(doc, objective)

add_heading(doc, '4. 課題に対する手法とアイデア', 1)
add_heading(doc, '4.1 基本方針', 2)
add_paragraph(doc, principle_p1)
add_paragraph(doc, principle_p2)
add_heading(doc, '4.2 全体処理の流れ', 2)
add_image(doc, FIG / 'overall_workflow.png', 16.0)
add_caption(doc, '図1　本パッケージの全体処理フロー')
add_paragraph(doc, method_p1)
add_heading(doc, '4.3 利用する外部情報源と役割分担', 2)
add_table(doc, ['情報源', '役割', '主な取得内容', '位置づけ'], source_rows, [4.0, 3.2, 3.3, 5.2])
add_heading(doc, '4.4 状態生成', 2)
add_paragraph(doc, state_gen_p1)
add_paragraph(doc, state_gen_p2)
add_heading(doc, '4.5 反応生成', 2)
add_image(doc, FIG / 'build_workflow.png', 16.0)
add_caption(doc, '図2　状態生成・反応生成・枝刈りの内部流れ')
add_paragraph(doc, reaction_gen_p1)
add_paragraph(doc, reaction_gen_p2)
add_heading(doc, '4.6 質量・電荷保存による欠損生成物の補完', 2)
add_paragraph(doc, balance_p)
add_heading(doc, '4.7 熱化学にもとづく枝刈り', 2)
add_paragraph(doc, thermo_p1)
add_paragraph(doc, thermo_p2)
add_heading(doc, '4.8 信頼度付与', 2)
add_paragraph(doc, confidence_p)
add_heading(doc, '4.9 自動収集・更新と再現性の確保', 2)
add_image(doc, FIG / 'ops_workflow.png', 16.0)
add_caption(doc, '図3　自動収集と再現性確保の運用フロー')
add_paragraph(doc, ops_p)
add_heading(doc, '4.10 可視化', 2)
add_paragraph(doc, visual_p1)
add_paragraph(doc, visual_p2)

add_heading(doc, '5. コード構成と工夫', 1)
add_paragraph(doc, code_p1)
add_table(doc, ['主なファイル', '役割'], module_rows, [5.1, 10.7])
add_paragraph(doc, extension_p)

add_heading(doc, '6. 手法による効果', 1)
add_heading(doc, '6.1 機能の有無による比較', 2)
add_paragraph(doc, ablation_p1)
add_image(doc, FIG / 'ablation_chart.png', 16.0)
add_caption(doc, '図4　主要機能の有無による出力件数の比較')
add_table(doc, ['条件', '状態数', '反応数', '読み取り'], [[a, str(b), str(c), d] for a,b,c,d in ablation_rows], [4.0, 1.8, 1.8, 8.2])
add_paragraph(doc, ablation_p2)
add_heading(doc, '6.2 現行版の生成結果の概況', 2)
add_paragraph(doc, results_p1)
add_paragraph(doc, results_p2)
add_image(doc, FIG / 'output_summary.png', 16.0)
add_caption(doc, '図5　サンプル生成結果の概況')

add_heading(doc, '7. 結果の妥当性と考察', 1)
add_paragraph(doc, validity_p1)
add_paragraph(doc, validity_p2)
add_paragraph(doc, validity_p3)
add_heading(doc, '7.1 例題から確認できる具体的事実', 2)
add_table(doc, ['項目', '値'], [
    ['投入ガス数', str(metadata['feed_count'])],
    ['生成状態数', str(metadata['species_count'])],
    ['生成反応数', str(metadata['reaction_count'])],
    ['内蔵状態辞書の件数', str(len(species_lib['species']))],
    ['内蔵反応雛型の件数', str(len(rx_ch4['reactions']) + len(rx_c4f8['reactions']))],
    ['原子準位拡張で追加された状態数', str(asd_added)],
    ['外部根拠情報から追加された反応雛型数', str(ext_added)],
    ['質量・電荷保存で補完された反応数', str(len(balance_completions))],
    ['熱化学枝刈りで除去された反応数', str(pruned_count)],
    ['可視化成果物数', str(visual_manifest['artifact_count'])],
], [9.5, 5.0])
add_heading(doc, '7.2 生成図の例', 2)
add_paragraph(doc, '以下の図は、本パッケージが実際に出力する工程俯瞰図と二部グラフの例である。図中の英語表記はコード出力そのままであるが、監査対象としての可視化例を示すために掲載する。', first_line_cm=0.0)
add_image(doc, FIG / 'engineer_process_dag_sample.png', 16.3)
add_caption(doc, '付図1　工程俯瞰図の出力例')
add_image(doc, FIG / 'plasma_bipartite_dag_sample.png', 16.3)
add_caption(doc, '付図2　二部グラフの出力例')

add_heading(doc, '8. 将来拡張性', 1)
add_table(doc, ['拡張項目', '内容'], future_rows, [4.3, 10.5])

add_heading(doc, '9. まとめ', 1)
add_paragraph(doc, summary_p)

add_heading(doc, '参考文献', 1)
for rid, text in refs:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(0.0)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(f'{rid} {text}')
    set_east_asian_font(run)
    run.font.size = Pt(10.0)

add_heading(doc, '付記', 1)
add_paragraph(doc, '編集用の Mermaid 記述は Markdown 版の付録に収録した。報告書本文は読みやすさを優先し、本文中では画像化した処理フローを用いた。', first_line_cm=0.0)

docx_path = OUT / 'plasma_reaction_builder_v10_report_revised_ja.docx'
doc.save(str(docx_path))

# Bundle zip
import zipfile
zip_path = Path('/mnt/data/plasma_reaction_builder_v10_report_revised_bundle.zip')
with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
    for path in OUT.rglob('*'):
        if path.is_file():
            zf.write(path, arcname=path.relative_to(OUT.parent))

print('Wrote', docx_path)
print('Wrote', OUT / 'plasma_reaction_builder_v10_report_revised_ja.md')
print('Wrote', zip_path)
